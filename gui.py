from tkinter import *
from tkinter import filedialog
from tkinter import messagebox
from os import path
import os
import pandas as pd
from docx.api import Document

window = Tk()
 
window.title("Pdf Parser")
#design interface
lbl = Label(window, text="Directory Path")
lbl.grid(column=0, row=0)
keys = ['Assignment name', 'Approx. value of the contract (in current US$)', 'Country', 'Duration of assignment (months)', 'Name of Client', 'Total No. of staff-months of the assignment', 'No./Address', 'Name', 'Title', 'Address', 'Tel', 'Email', 'Start date (month/year)', 'Completion date (month/year)', 'No. of professional staff-months provided by your consulting firm/organization or your sub consultants', 'Name of associated Consultants, if any', 'Name of senior professional staff of your consulting firm/organization involved and designation and/or functions performed (e.g. Project Director/Coordinator, Team Leader)', 'Description of Project', 'Description of actual services provided by your staff within the assignment']
ent_path = Entry(window,width=50)
ent_path.grid(column=1, row=0)

def iter_unique_cells(row):
    """Generate cells in *row* skipping empty grid cells."""
    prior_tc = None
    for cell in row.cells:
        this_tc = cell._tc
        if this_tc is prior_tc:
            continue
        prior_tc = this_tc
        yield cell
        
def parse():
    data = []
    dirpath = ent_path.get()
    for path in os.listdir():
        full_path = os.path.join(dirpath, path)
        if os.path.isfile(full_path) and path.endswith(".docx"):
            document = Document(full_path)
            if len(document.tables) == 0:
                continue
            table = document.tables[0]
            data.append([])
            flag = False
            for row in table.rows:
                for cell in iter_unique_cells(row):
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if not len(text):
                            continue
                        if text[-1] == ':':
                            data[-1].append('')
                            flag = True
                        else:
                            texts = text.split(':')
                            if len(texts) == 1 and flag:
                                data[-1][-1] += text
                            elif len(texts) > 1:
                                data[-1].append(texts[1])
                                flag = False
    df = pd.DataFrame(data, columns = keys)
    df.to_excel("output.xlsx", index=False)                   


def browser_clicked():
    file_tmp = filedialog.askdirectory(initialdir= path.dirname(__file__))
    ent_path.delete(0,END)
    ent_path.insert(0,file_tmp)

btn_browser = Button(window, text="Browser", command=browser_clicked)
btn_browser.grid(column=2, row=0)

btn_parse = Button(window, text="Parse", command=parse)
btn_parse.grid(column=1, row=1)

window.mainloop()
