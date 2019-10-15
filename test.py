import pandas as pd
from docx.api import Document

document = Document('1.1.docx')
table = document.tables[0]

data = [[]]
keys = ['Assignment name', 'Approx. value of the contract (in current US$)', 'Country', 'Duration of assignment (months)', 'Name of Client', 'Total No. of staff-months of the assignment', 'No./Address', 'Name', 'Title', 'Address', 'Tel', 'Email', 'Start date (month/year)', 'Completion date (month/year)', 'No. of professional staff-months provided by your consulting firm/organization or your sub consultants', 'Name of associated Consultants, if any', 'Name of senior professional staff of your consulting firm/organization involved and designation and/or functions performed (e.g. Project Director/Coordinator, Team Leader)', 'Description of Project', 'Description of actual services provided by your staff within the assignment']

def iter_unique_cells(row):
    """Generate cells in *row* skipping empty grid cells."""
    prior_tc = None
    for cell in row.cells:
        this_tc = cell._tc
        if this_tc is prior_tc:
            continue
        prior_tc = this_tc
        yield cell
i = 0
flag = False
for row in table.rows:
    for cell in iter_unique_cells(row):
        for para in cell.paragraphs:
            text = para.text.strip()
            if not len(text):
                continue
            if text[-1] == ':':
                data[0].append('')
                flag = True
            else:
                texts = text.split(':')
                if len(texts) == 1 and flag:
                    data[0][-1] += text
                elif len(texts) > 1:
                    data[0].append(texts[1])
                    flag = False
# df = pd.DataFrame(data, columns = keys)
# df.to_excel("output.xlsx", index=False)
#print(data)

print (len(data[0]))
for d in data[0]:
    print (d)
    print("-----------------")